from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

def generate_txt(content, path):
    with open(path,"w") as f:
        f.write(content)

def generate_docx(content, path):
    doc=Document()
    doc.add_heading("AI Image Scan Report", level=1)
    doc.add_paragraph(content)
    doc.save(path)

def generate_pdf(content, path):
    c=canvas.Canvas(path, pagesize=A4)
    text=c.beginText(40, 800)
    for line in content.splitlines():
        text.textLine(line)
    c.drawText(text)
    c.save()
